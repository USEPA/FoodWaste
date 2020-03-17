# qar5_docx.py (qar5)
# !/usr/bin/env python3
# coding=utf-8
# young.daniel@epa.gov

"""Definition of qar5 docx export methods."""

from docx import Document
from docx.enum.style import WD_BUILTIN_STYLE, WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_COLOR_INDEX, WD_LINE_SPACING, \
    WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

from io import BytesIO
from os import path
import tempfile
from zipfile import ZipFile
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect
from django.templatetags.static import static
from DataSearch.settings import DEBUG, STATIC_ROOT
from qar5.views import get_qar5_for_user, get_qar5_for_team, get_qapp_info

def export_doc(request, *args, **kwargs):
    """Function to export a multiple QAR5 objects as Word Docx files."""
    if 'user' in request.path:
        user_id = kwargs.get('pk', None)
        team_id = qapp_id = None
    elif 'team' in request.path:
        team_id = kwargs.get('pk', None)
        user_id = qapp_id = None
    else:
        qapp_id = kwargs.get('pk', None)
        team_id = user_id = None

    if qapp_id is None or user_id or team_id:
        if user_id:
            qapp_ids = get_qar5_for_user(
                user_id).values_list('id', flat=True)
        else:
            qapp_ids = get_qar5_for_team(
                team_id).values_list('id', flat=True)

        zip_mem = BytesIO()
        archive = ZipFile(zip_mem, 'w')
        for id in qapp_ids:
            resp = export_doc_single(request, pk=id)
            filename = resp['filename']
            if filename:
                temp_file_name = '%d_%s' % (id, filename)
                with tempfile.SpooledTemporaryFile() as tmp:
                    archive.writestr(temp_file_name, resp.content)

        archive.close()
        response = HttpResponse(
            zip_mem.getvalue(), content_type='application/force-download')
        response['Content-Disposition'] = \
            'attachment; filename="%s_qapps.zip"' % request.user.username
        response['Content-length'] = zip_mem.tell()
        return response


def add_center_heading(document, text, level):
    """Helper method to easily add centered headers to a docx file"""
    heading_style = 'Heading %d' % level
    paragraph = document.add_paragraph(text, heading_style)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


#def add_center_heading_styles(document):
#    """
#    Helper method to instantiate a document with the required header formats
#    """
#    style_center = document.styles.add_style(
#        'center_text', WD_STYLE_TYPE.PARAGRAPH).paragraph_format
#    style_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def set_table_row_height(table):
    """Helper method to set minimum row height and alignment for a table"""
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.35)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER



# py-lint: disable=no-member
def export_doc_single(request, *args, **kwargs):
    """Function to export a single QAR5 object as a Word Docx file."""
    qapp_id = kwargs.get('pk', None) 
    qapp_info = get_qapp_info(request.user, qapp_id)

    if not qapp_info:
        return HttpResponseRedirect(request)

    filename = '%s.docx' % qapp_info['qapp'].title

    document = Document()
    styles = document.styles
        
    # #################################################
    # BEGIN COVER PAGE
    # #################################################
    # Coversheet with signatures section:
    if DEBUG:
        logo = path.join(STATIC_ROOT, 'EPA_Files', 'logo.png')
        qual_assur_proj_plan = path.join(
            STATIC_ROOT, 'images', 'quality_assurance_project_plan.PNG')
    else:
        logo = static('logo.png')
        qual_assur_proj_plan = static('quality_assurance_project_plan.PNG')

    # TODO: Add top row logo and blue background label
    run = document.add_paragraph().add_run()
    run.add_picture(logo, width=Inches(1.5))
    run.add_text('\t\t\t')
    run.add_picture(qual_assur_proj_plan, width=Inches(3))

    style_center = document.styles.add_style(
        'center_text', WD_STYLE_TYPE.PARAGRAPH).paragraph_format
    style_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    style_paragraph = document.styles.add_style(
        'style_paragraph', WD_STYLE_TYPE.PARAGRAPH).paragraph_format
    style_paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # TODO Make blue_header text white, add blue background with shadow
    # document.add_picture('blue_background.png', width=Inches(4))
    # background color: rgb(0, 176, 240)

    # The rest of the document will be WD_ALIGN_PARAGRAPH.CENTER

    # blank line
    add_center_heading(document, 'Office of Research and Development', 1)
    add_center_heading(document, qapp_info['qapp'].division.name, 1)
    # blank line
    # Next few sections are from the qapp object
    add_center_heading(document, qapp_info['qapp'].division_branch, level=3)
    # blank line
    add_center_heading(document, 'EPA Project Lead', level=2)
    for lead in qapp_info['qapp_leads']:
        add_center_heading(document, lead.name, level=3)
    # blank line
    add_center_heading(document, qapp_info['qapp'].intra_extra, level=3)
    add_center_heading(document, qapp_info['qapp'].qa_category, level=3)
    add_center_heading(document, qapp_info['qapp'].revision_number, level=3)
    add_center_heading(document, str(qapp_info['qapp'].date), level=3)
    # blank line
    add_center_heading(document, 'Prepared By', level=2)
    add_center_heading(document, 
        '%s %s' % (qapp_info['qapp'].prepared_by.first_name,
                    qapp_info['qapp'].prepared_by.last_name,),
        level=3)
    # blank line
    add_center_heading(document, qapp_info['qapp'].strap, level=3)
    add_center_heading(document, qapp_info['qapp'].tracking_id, level=3)

    # #################################################
    # END COVER PAGE
    # BEGIN APPROVAL PAGE
    # #################################################

    document.add_heading('Approval Page', level=2)
    # Signature grid ...
    num_signatures = len(qapp_info['signatures'])
    table = document.add_table(rows=6+num_signatures, cols=12)
    table.style = styles['Table Grid']

    set_table_row_height(table)

    row_cells = table.rows[0].cells
    row_cells[0].text = 'QA Project Plan Title:'
    row_cells[0].merge(row_cells[3])
    row_cells[4].text = qapp_info['qapp_approval'].project_plan_title
    row_cells[4].merge(row_cells[11])

    row_cells = table.rows[1].cells
    row_cells[0].text = 'QA Activity Number:'
    row_cells[0].merge(row_cells[3])
    row_cells[4].text = qapp_info['qapp_approval'].activity_number
    row_cells[4].merge(row_cells[11])

    # TODO: Center text in this row:
    row_cells = table.rows[2].cells
    row_cells[0].text = 'If Intramural or Extramural, EPA Project Approvals'
    row_cells[0].merge(row_cells[11])

    iter_count = 0
    # TODO: Iterate through EPA Project Approvals:
    # Start with row 3 + iter_count++
    for sig in qapp_info['signatures']:
        if not sig.contractor:
            row_cells = table.rows[3 + iter_count].cells
            row_cells[0].text = 'Name:'
            row_cells[1].text = sig.name
            row_cells[1].merge(row_cells[3])

            row_cells[4].text = 'Signature/Date:'
            row_cells[4].merge(row_cells[5])
            row_cells[6].merge(row_cells[11])
            iter_count += 1

    # Always insert a blank entry for hand-written approval sigs
    row_cells = table.rows[3 + iter_count].cells
    row_cells[0].text = 'Name:'
    row_cells[1].merge(row_cells[3])
    row_cells[4].text = 'Signature/Date:'
    row_cells[4].merge(row_cells[5])
    row_cells[6].merge(row_cells[11])

    # TODO: Center text in this row:
    row_cells = table.rows[4 + iter_count].cells
    row_cells[0].text = 'If Extramural, Contractor Project Approvals'
    row_cells[0].merge(row_cells[11])

    # TODO: Iterate through Contractor Project Approvals:
    # Start with row 5 + iter_count++
    for sig in qapp_info['signatures']:
        if sig.contractor:
            row_cells = table.rows[5 + iter_count].cells
            row_cells[0].text = 'Name:'
            row_cells[1].text = sig.name
            row_cells[1].merge(row_cells[3])

            row_cells[4].text = 'Signature/Date:'
            row_cells[4].merge(row_cells[5])
            row_cells[6].merge(row_cells[11])
            iter_count += 1

    # Always insert a blank entry for hand-written approval sigs
    row_cells = table.rows[5 + iter_count].cells
    row_cells[0].text = 'Name:'
    row_cells[1].merge(row_cells[3])
    row_cells[4].text = 'Signature/Date:'
    row_cells[4].merge(row_cells[5])
    row_cells[6].merge(row_cells[11])
    document.add_page_break()

    # #################################################
    # END APPROVAL PAGE
    # BEGIN ToC PAGE
    # #################################################

    document.add_heading('Table of Contents', level=2)
    document.add_heading(
        'TODO: This will have to be generated after the rest of ' + \
        'the doc so we know page numbers and contents...', level=3)
    document.add_page_break()

    # #################################################
    # END ToC PAGE
    # BEGIN Everything Else PAGE
    # #################################################

    #  1) Heading 1 - Revision History
    document.add_heading('Revision History', level=1)
    #  2) Table Label
    document.add_heading('Table 1 QAR5 Revision History', level=3)
    #  3) Table (revision history)
    num_revisions = len(qapp_info['revisions'])
    table = document.add_table(rows=1+num_signatures, cols=3)
    table.style = styles['Light List']
    row_cells = table.rows[0].cells
    row_cells[0].text = 'Revision Number'
    row_cells[1].text = 'Date Approved'
    row_cells[2].text = 'Revision'

    iter_count = 0
    for rev in qapp_info['revisions']:
        row_cells = table.rows[1+iter_count].cells
        row_cells[0].text = rev.revision
        row_cells[1].text = str(rev.effective_date)
        row_cells[2].text = rev.description
        iter_count += 1

    # TODO: Paragraphs aren't formatting properly, still double spaces...

    # Section A
    document.add_heading('Section A - Executive Summary', level=1)
    if qapp_info['section_a']:
        document.add_heading('A.1 Distribution List', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a3.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('A.2 Project Task Organization', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a4.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('A.3 Problem Definition Background', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a5.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('A.4 Project Description', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a6.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('A.5 Quality Objectives and Criteria', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a7.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('A.6 Special Training Certification', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a8.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('A.7 Documents and Records', level=2)
        document.add_paragraph(
            qapp_info['section_a'].a9.replace('\r\n', ' '),
            styles['No Spacing'])
    else:
        document.add_heading('SECTION A INCOMPLETE!', level=2)

    # Section B
    document.add_heading('Section B - Experimental Design', level=1)
    if qapp_info['section_b']:
        document.add_heading(
            'B.1 Sample/Data Collection, Gathering, or Use', level=2)
        document.add_heading('B.1.1 Use', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b1_2.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.1.2 Requirements', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b1_3.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.1.3 Databases, Maps, Literature', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b1_4.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.1.4 Non-Quality Constraints', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b1_5.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading(
            'B.2 Data Analysis / Statistical Design / Data Management',
            level=2)
        document.add_heading('B.2.1 Sources', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b2_1.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.2.2 Acceptance/Rejection Process', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b2_2.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.2.3 Rationale for Selections', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b2_3.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.2.4 Procedures', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b2_4.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.2.5 Disclaimer', level=3)
        document.add_paragraph(
            qapp_info['section_b'].b2_5.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.3 Data Management and Documentation', level=2)
        document.add_paragraph(
            qapp_info['section_b'].b3.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('B.4 Tracking', level=2)
        document.add_paragraph(
            qapp_info['section_b'].b4.replace('\r\n', ' '),
            styles['No Spacing'])
    else:
        document.add_heading('SECTION B INCOMPLETE!', level=2)

    # Section C
    document.add_heading('Section C', level=1)
    if qapp_info['section_c']:
        document.add_heading(
            'C.1 Assessments and Response Actions', level=2)
        document.add_paragraph(
            qapp_info['section_c'].c1.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading('C.2 Reports to Management', level=2)
        document.add_paragraph(
            qapp_info['section_c'].c2.replace('\r\n', ' '),
            styles['No Spacing'])
    else:
        document.add_heading(
            'C.1 Assessments and Response Actions', level=2)
        document.add_heading('C.2 Reports to Management', level=2)

    # Section D
    if qapp_info['section_d']:
        document.add_heading('Section D', level=1)
        document.add_heading(
            'D.1 Data Review, Verification, and Validation', level=2)
        document.add_paragraph(
            qapp_info['section_d'].d1.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading(
            'D.2 Verification and Validation Methods', level=2)
        document.add_paragraph(
            qapp_info['section_d'].d2.replace('\r\n', ' '),
            styles['No Spacing'])
        document.add_heading(
            'D.3 Reconciliation with User Requirements', level=2)
        document.add_paragraph(
            qapp_info['section_d'].d3.replace('\r\n', ' '),
            styles['No Spacing'])
    else:
        document.add_heading('SECTION D INCOMPLETE!', level=2)

    # References
    document.add_heading('References', level=1)
    if qapp_info['references']:
        document.add_paragraph(
            qapp_info['references'].references.replace('\r\n\r\n', '\r\n'),
            styles['No Spacing'])
    else:
        document.add_heading('REFERENCES SECTION INCOMPLETE!', level=2)

    content_type='application/vnd.openxmlformats-officedocument.' + \
            'wordprocessingml.document'
    response = HttpResponse(content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % filename
    document.save(response)
    response['filename'] = filename
    return response
